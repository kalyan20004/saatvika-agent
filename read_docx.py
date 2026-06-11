import docx

doc = docx.Document('E:\\microsoft-buildAI\\Saatvika_TestCases.docx')
with open('E:\\microsoft-buildAI\\Saatvika_TestCases.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text + '\n')

    f.write("\n--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            f.write(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]) + "\n")
        f.write("-" * 40 + "\n")
